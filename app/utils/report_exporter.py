"""
æŠ¥å‘Šå¯¼å‡ºå·¥å…· - æ”¯æŒ Markdownã€Wordã€PDF æ ¼å¼

ä¾èµ–å®‰è£…:
    pip install pypandoc markdown

PDF å¯¼å‡ºéœ€è¦é¢å¤–å·¥å…·:
    - wkhtmltopdf (æ¨è): https://wkhtmltopdf.org/downloads.html
    - æˆ– LaTeX: https://www.latex-project.org/get/
"""

import logging
import os
import tempfile
from pathlib import Path
from typing import Dict, Any, Optional

logger = logging.getLogger(__name__)

# æ£€æŸ¥ä¾èµ–æ˜¯å¦å¯ç”¨
try:
    import markdown
    import pypandoc

    # æ£€æŸ¥ pandoc æ˜¯å¦å¯ç”¨
    try:
        pypandoc.get_pandoc_version()
        PANDOC_AVAILABLE = True
        logger.info("âœ… Pandoc å¯ç”¨")
    except OSError:
        PANDOC_AVAILABLE = False
        logger.warning("âš ï¸ Pandoc ä¸å¯ç”¨ï¼ŒWord å’Œ PDF å¯¼å‡ºåŠŸèƒ½å°†ä¸å¯ç”¨")

    EXPORT_AVAILABLE = True
except ImportError as e:
    EXPORT_AVAILABLE = False
    PANDOC_AVAILABLE = False
    logger.warning(f"âš ï¸ å¯¼å‡ºåŠŸèƒ½ä¾èµ–åŒ…ç¼ºå¤±: {e}")
    logger.info("ğŸ’¡ è¯·å®‰è£…: pip install pypandoc markdown")

# æ£€æŸ¥ pdfkitï¼ˆå”¯ä¸€çš„ PDF ç”Ÿæˆå·¥å…·ï¼‰
PDFKIT_AVAILABLE = False
PDFKIT_ERROR = None

try:
    import pdfkit
    # æ£€æŸ¥ wkhtmltopdf æ˜¯å¦å®‰è£…
    try:
        pdfkit.configuration()
        PDFKIT_AVAILABLE = True
        logger.info("âœ… pdfkit + wkhtmltopdf å¯ç”¨ï¼ˆPDF ç”Ÿæˆå·¥å…·ï¼‰")
    except Exception as e:
        PDFKIT_ERROR = str(e)
        logger.warning("âš ï¸ wkhtmltopdf æœªå®‰è£…ï¼ŒPDF å¯¼å‡ºåŠŸèƒ½ä¸å¯ç”¨")
        logger.info("ğŸ’¡ å®‰è£…æ–¹æ³•: https://wkhtmltopdf.org/downloads.html")
except ImportError:
    logger.warning("âš ï¸ pdfkit æœªå®‰è£…ï¼ŒPDF å¯¼å‡ºåŠŸèƒ½ä¸å¯ç”¨")
    logger.info("ğŸ’¡ å®‰è£…æ–¹æ³•: pip install pdfkit")
except Exception as e:
    PDFKIT_ERROR = str(e)
    logger.warning(f"âš ï¸ pdfkit æ£€æµ‹å¤±è´¥: {e}")

# æ£€æŸ¥ weasyprint (æ›¿ä»£ PDF ç”Ÿæˆå·¥å…·)
WEASYPRINT_AVAILABLE = False
WEASYPRINT_ERROR = None

try:
    import weasyprint
    WEASYPRINT_AVAILABLE = True
    logger.info("âœ… weasyprint å¯ç”¨")
except ImportError:
    logger.warning("âš ï¸ weasyprint æœªå®‰è£…")
    logger.info("ğŸ’¡ å®‰è£…æ–¹æ³•: brew install pango && pip install weasyprint")
except Exception as e:
    WEASYPRINT_ERROR = str(e)
    logger.warning(f"âš ï¸ weasyprint æ£€æµ‹å¤±è´¥: {e}")



class ReportExporter:
    """æŠ¥å‘Šå¯¼å‡ºå™¨ - æ”¯æŒ Markdownã€Wordã€PDF æ ¼å¼"""

    def __init__(self):
        self.export_available = EXPORT_AVAILABLE
        self.pandoc_available = PANDOC_AVAILABLE
        self.pdfkit_available = PDFKIT_AVAILABLE
        self.weasyprint_available = WEASYPRINT_AVAILABLE

        logger.info("ğŸ“‹ ReportExporter åˆå§‹åŒ–:")
        logger.info(f"  - export_available: {self.export_available}")
        logger.info(f"  - pandoc_available: {self.pandoc_available}")
        logger.info(f"  - pdfkit_available: {self.pdfkit_available}")
        logger.info(f"  - weasyprint_available: {self.weasyprint_available}")
    
    def _format_json_content(self, key: str, content: str) -> str:
        """å°è¯•è§£æå¹¶æ ¼å¼åŒ–JSONå†…å®¹"""
        try:
            import json
            import ast
            
            content = content.strip()
            if not content.startswith('{'):
                return content
                
            try:
                data = json.loads(content)
            except json.JSONDecodeError:
                # å°è¯•å¤„ç† Python å­—å…¸å­—ç¬¦ä¸² (å•å¼•å·)
                try:
                    data = ast.literal_eval(content)
                    if not isinstance(data, dict):
                        return content
                except:
                    return content
            
            if not isinstance(data, dict):
                return content
                
            # æ ¹æ®ä¸åŒçš„æ¨¡å—keyè¿›è¡Œç‰¹å®šæ ¼å¼åŒ–
            if key == "macro_report":
                return self._format_macro_report(data)
            elif key == "policy_report":
                return self._format_policy_report(data)
            elif key == "sector_report":
                return self._format_sector_report(data)
            elif key == "international_news_report":
                return self._format_international_news_report(data)
            elif key == "strategy_report" or key == "research_team_decision":
                return self._format_strategy_report(data)
            elif key == "technical_report":
                return self._format_technical_report(data)
            else:
                return content # å…¶ä»–æ¨¡å—æš‚ä¸å¤„ç†ï¼Œä¿æŒåŸæ ·æˆ–é€šç”¨æ ¼å¼åŒ–
                
        except Exception:
            # è§£æå¤±è´¥æˆ–éJSONå†…å®¹ï¼Œè¿”å›åŸå§‹å†…å®¹
            return content

    def _format_macro_report(self, data: Dict[str, Any]) -> str:
        lines = []
        lines.append("### ğŸ“Š æ ¸å¿ƒè§‚ç‚¹")
        if "economic_cycle" in data:
            lines.append(f"- **ç»æµå‘¨æœŸ**: {data['economic_cycle']}")
        if "liquidity" in data:
            lines.append(f"- **æµåŠ¨æ€§ç¯å¢ƒ**: {data['liquidity']}")
        if "sentiment_score" in data:
            lines.append(f"- **æƒ…ç»ªè¯„åˆ†**: {data['sentiment_score']}")
        if "confidence" in data:
            lines.append(f"- **ç½®ä¿¡åº¦**: {data['confidence']}")
        lines.append("")
        
        if "key_indicators" in data and data["key_indicators"]:
            lines.append("### ğŸ“ˆ å…³é”®æŒ‡æ ‡")
            for indicator in data["key_indicators"]:
                lines.append(f"- {indicator}")
            lines.append("")
            
        if "analysis_summary" in data:
            lines.append("### ğŸ“ åˆ†ææ€»ç»“")
            lines.append(data["analysis_summary"])
            lines.append("")
            
        if "data_note" in data:
            lines.append(f"> âš ï¸ {data['data_note']}")
            
        return "\n".join(lines)

    def _format_policy_report(self, data: Dict[str, Any]) -> str:
        lines = []
        lines.append("### ğŸ“œ æ”¿ç­–ç¯å¢ƒ")
        # å…¼å®¹ä¸åŒå­—æ®µå
        support = data.get("support_strength") or data.get("overall_support_strength")
        if support:
            lines.append(f"- **æ”¯æŒåŠ›åº¦**: {support}")
            
        continuity = data.get("policy_continuity")
        if continuity:
            lines.append(f"- **æ”¿ç­–è¿ç»­æ€§**: {continuity}")
            
        if "confidence" in data:
            lines.append(f"- **ç½®ä¿¡åº¦**: {data['confidence']}")
        lines.append("")
        
        # å…¼å®¹ key_policies å’Œ key_events
        policies = data.get("key_policies") or data.get("key_events")
        if policies:
            lines.append("### ğŸ—ï¸ å…³é”®æ”¿ç­–")
            for policy in policies:
                lines.append(f"- {policy}")
            lines.append("")
            
        if "industry_policy" in data and data["industry_policy"]:
            lines.append("### ğŸ­ äº§ä¸šæ”¿ç­–")
            for policy in data["industry_policy"]:
                lines.append(f"- {policy}")
            lines.append("")

        if "long_term_policies" in data and data["long_term_policies"]:
            lines.append("### ğŸ”­ é•¿æœŸæˆ˜ç•¥")
            for policy in data["long_term_policies"]:
                if isinstance(policy, dict):
                    name = policy.get("name", "")
                    duration = policy.get("duration", "")
                    lines.append(f"- **{name}** ({duration})")
                else:
                    lines.append(f"- {policy}")
            lines.append("")
            
        if "analysis_summary" in data:
            lines.append("### ğŸ“ åˆ†ææ€»ç»“")
            lines.append(data["analysis_summary"])
            lines.append("")
            
        if "impact_sector" in data and data["impact_sector"]:
            lines.append("### ğŸ¯ å½±å“æ¿å—")
            lines.append(f"{', '.join(data['impact_sector'])}")
            
        return "\n".join(lines)

    def _format_sector_report(self, data: Dict[str, Any]) -> str:
        lines = []
        # å…¼å®¹ market_style å’Œ rotation_trend
        style = data.get("market_style") or data.get("rotation_trend")
        if style:
            lines.append(f"**å¸‚åœºé£æ ¼/è½®åŠ¨**: {style}")
            
        if "heat_score" in data:
            lines.append(f"**çƒ­åº¦è¯„åˆ†**: {data['heat_score']}")
        lines.append("")
        
        if "hot_themes" in data and data["hot_themes"]:
            lines.append("### ğŸ”¥ çƒ­é—¨ä¸»é¢˜")
            for theme in data["hot_themes"]:
                lines.append(f"- {theme}")
            lines.append("")
            
        if "top_sectors" in data and data["top_sectors"]:
            lines.append("### ğŸš€ é¢†æ¶¨æ¿å—")
            for sector in data["top_sectors"]:
                lines.append(f"- {sector}")
            lines.append("")
            
        if "bottom_sectors" in data and data["bottom_sectors"]:
            lines.append("### ğŸ“‰ é¢†è·Œæ¿å—")
            for sector in data["bottom_sectors"]:
                lines.append(f"- {sector}")
            lines.append("")
            
        if "analysis_summary" in data:
            lines.append("### ğŸ“ æ¿å—é€»è¾‘")
            lines.append(data["analysis_summary"])
            
        return "\n".join(lines)

    def _format_international_news_report(self, data: Dict[str, Any]) -> str:
        lines = []
        lines.append("### ğŸŒ å›½é™…æ–°é—»å½±å“")
        
        # å…¼å®¹ impact_strength
        strength = data.get("impact_strength") or data.get("overall_impact")
        if strength:
            lines.append(f"- **å½±å“å¼ºåº¦**: {strength}")
            
        if "impact_duration" in data:
            lines.append(f"- **å½±å“æŒç»­æ€§**: {data['impact_duration']}")
        if "risk_level" in data:
            lines.append(f"- **é£é™©ç­‰çº§**: {data['risk_level']}")
        lines.append("")
        
        # å…¼å®¹ key_events å’Œ key_news
        events = data.get("key_events") or data.get("key_news")
        if events:
            lines.append("### ğŸ“° å…³é”®äº‹ä»¶")
            for event in events:
                if isinstance(event, dict):
                    # å¤„ç†å¯èƒ½çš„å­—å…¸ç»“æ„ (title, summaryç­‰)
                    title = event.get("title", "")
                    summary = event.get("summary", "")
                    lines.append(f"- **{title}**: {summary}")
                else:
                    lines.append(f"- {event}")
            lines.append("")
            
        if "analysis_summary" in data:
            lines.append("### ğŸ“ åˆ†ææ€»ç»“")
            lines.append(data["analysis_summary"])
            
        return "\n".join(lines)

    def _format_technical_report(self, data: Dict[str, Any]) -> str:
        lines = []
        if "trend_signal" in data:
            lines.append(f"**è¶‹åŠ¿ä¿¡å·**: {data['trend_signal']}")
        if "confidence" in data:
            lines.append(f"**ç½®ä¿¡åº¦**: {data['confidence']}")
        lines.append("")
        
        if "key_levels" in data:
            lines.append("### ğŸ¯ å…³é”®ç‚¹ä½")
            levels = data["key_levels"]
            if isinstance(levels, dict):
                for k, v in levels.items():
                    lines.append(f"- **{k}**: {v}")
            elif isinstance(levels, list):
                for l in levels:
                    lines.append(f"- {l}")
            lines.append("")
            
        if "indicators" in data:
            lines.append("### ğŸ“ˆ æŠ€æœ¯æŒ‡æ ‡")
            indicators = data["indicators"]
            if isinstance(indicators, dict):
                for k, v in indicators.items():
                    lines.append(f"- **{k}**: {v}")
            lines.append("")
            
        if "analysis_summary" in data:
            lines.append("### ğŸ“ æŠ€æœ¯åˆ†æ")
            lines.append(data["analysis_summary"])
            
        return "\n".join(lines)

    def _format_strategy_report(self, data: Dict[str, Any]) -> str:
        # å¤„ç†åµŒå¥—çš„ strategy_report (å½“ä¼ å…¥çš„æ˜¯æ•´ä¸ªèŠ‚ç‚¹è¾“å‡ºæ—¶)
        if "strategy_report" in data:
            inner = data["strategy_report"]
            if isinstance(inner, str):
                try:
                    import json
                    inner = json.loads(inner)
                except:
                    pass
            if isinstance(inner, dict):
                data = inner

        lines = []
        if "market_outlook" in data:
            lines.append(f"### ğŸ¯ å¸‚åœºå±•æœ›: {data['market_outlook']}")
        
        if "final_position" in data:
            pos = data["final_position"]
            if isinstance(pos, (int, float)):
                pos = f"{pos:.2%}"
            lines.append(f"### ğŸ’¼ å»ºè®®ä»“ä½: {pos}")
        lines.append("")
        
        if "position_breakdown" in data:
            pb = data["position_breakdown"]
            lines.append("#### ğŸ—ï¸ ä»“ä½ç»“æ„")
            if "core_holding" in pb:
                lines.append(f"- **æ ¸å¿ƒé•¿æœŸä»“ä½**: {pb['core_holding']:.2%}")
            if "tactical_allocation" in pb:
                lines.append(f"- **æˆ˜æœ¯é…ç½®**: {pb['tactical_allocation']:.2%}")
            if "cash_reserve" in pb:
                lines.append(f"- **ç°é‡‘å‚¨å¤‡**: {pb['cash_reserve']:.2%}")
            lines.append("")
            
        if "adjustment_triggers" in data:
            at = data["adjustment_triggers"]
            lines.append("#### ğŸ”” åŠ¨æ€è°ƒæ•´è§¦å‘")
            if "increase_to" in at and "increase_condition" in at:
                lines.append(f"- ğŸ“ˆ **åŠ ä»“è‡³ {at['increase_to']:.2%}**: {at['increase_condition']}")
            if "decrease_to" in at and "decrease_condition" in at:
                lines.append(f"- ğŸ“‰ **å‡ä»“è‡³ {at['decrease_to']:.2%}**: {at['decrease_condition']}")
            lines.append("")
            
        if "opportunity_sectors" in data and data["opportunity_sectors"]:
            lines.append("#### ğŸš€ æœºä¼šæ¿å—")
            for s in data["opportunity_sectors"]:
                lines.append(f"- {s}")
            lines.append("")
            
        if "key_risks" in data and data["key_risks"]:
            lines.append("#### âš ï¸ é£é™©æç¤º")
            for r in data["key_risks"]:
                lines.append(f"- {r}")
            lines.append("")
            
        if "debate_summary" in data and data["debate_summary"] and data["debate_summary"] != "æ— è¾©è®ºæ€»ç»“":
            lines.append("#### ğŸ—³ï¸ è¾©è®ºæ€»ç»“")
            summary = data["debate_summary"]
            # ç¡®ä¿æ®µè½æ­£ç¡®åˆ’åˆ†
            summary = summary.replace("\n", "\n\n")
            lines.append(summary)
            lines.append("")
        elif "current_response" in data and data["current_response"]:
            # å…¼å®¹å¤„ç†ï¼šå¦‚æœæ²¡æœ‰debate_summaryä½†æœ‰current_responseï¼ˆé€šå¸¸æ˜¯æŒ‡æ•°åˆ†æçš„æœ€åè¾©è®ºå›åˆï¼‰
            lines.append("#### ğŸ—³ï¸ è¾©è®ºç„¦ç‚¹")
            summary = data["current_response"]
            summary = summary.replace("\n", "\n\n")
            lines.append(summary)
            lines.append("")

        if "rationale" in data:
            lines.append("#### ğŸ’¡ ç­–ç•¥é€»è¾‘")
            rationale = data["rationale"]
            # ç¡®ä¿æ®µè½æ­£ç¡®åˆ’åˆ†
            rationale = rationale.replace("\n", "\n\n")
            lines.append(rationale)
            lines.append("")
            
        # ç§»é™¤å†³ç­–å…¬å¼ï¼Œé¿å…æŠ¥å‘Šå†—ä½™
        # if "decision_rationale" in data:
        #     lines.append(f"> ğŸ”¢ **å†³ç­–å…¬å¼**: {data['decision_rationale']}")
            
        return "\n".join(lines)

    def generate_markdown_report(self, report_doc: Dict[str, Any]) -> str:

        """ç”Ÿæˆ Markdown æ ¼å¼æŠ¥å‘Š"""
        logger.info("ğŸ“ ç”Ÿæˆ Markdown æŠ¥å‘Š...")
        
        stock_symbol = report_doc.get("stock_symbol", "unknown")
        analysis_date = report_doc.get("analysis_date", "")
        analysts = report_doc.get("analysts", [])
        research_depth = report_doc.get("research_depth", 1)
        reports = report_doc.get("reports", {})
        summary = report_doc.get("summary", "")
        
        content_parts = []
        
        # æ ‡é¢˜å’Œå…ƒä¿¡æ¯
        content_parts.append(f"# {stock_symbol} è‚¡ç¥¨åˆ†ææŠ¥å‘Š")
        content_parts.append("")
        content_parts.append(f"**åˆ†ææ—¥æœŸ**: {analysis_date}")
        if analysts:
            content_parts.append(f"**åˆ†æå¸ˆ**: {', '.join(analysts)}")
        content_parts.append(f"**ç ”ç©¶æ·±åº¦**: {research_depth}")
        content_parts.append("")
        content_parts.append("---")
        content_parts.append("")
        
        # æ‰§è¡Œæ‘˜è¦
        if summary:
            content_parts.append("## ğŸ“Š æ‰§è¡Œæ‘˜è¦")
            content_parts.append("")
            content_parts.append(summary)
            content_parts.append("")
            content_parts.append("---")
            content_parts.append("")
        
        # å„æ¨¡å—å†…å®¹
        module_order = [
            "company_overview",
            "macro_report",  # æŒ‡æ•°åˆ†æ
            "policy_report", # æŒ‡æ•°åˆ†æ
            "financial_analysis", 
            "sector_report", # æŒ‡æ•°åˆ†æ
            "technical_analysis",
            "technical_report", # æŒ‡æ•°åˆ†æ (åˆ«å)
            "market_analysis",
            "international_news_report", # æŒ‡æ•°åˆ†æ
            "risk_analysis",
            "valuation_analysis",
            "valuation_report", # æŒ‡æ•°åˆ†æ (åˆ«å)
            "sentiment_report", # æŒ‡æ•°åˆ†æ
            "market_sentiment", # æŒ‡æ•°åˆ†æ
            "investment_recommendation",
            "strategy_report"   # æŒ‡æ•°åˆ†æ
        ]
        
        module_titles = {
            "company_overview": "ğŸ¢ å…¬å¸æ¦‚å†µ",
            "macro_report": "ğŸŒ å®è§‚ç»æµåˆ†æ",
            "policy_report": "ğŸ“œ æ”¿ç­–ç¯å¢ƒåˆ†æ",
            "financial_analysis": "ğŸ’° è´¢åŠ¡åˆ†æ",
            "sector_report": "ğŸ™ï¸ è¡Œä¸šæ¿å—åˆ†æ",
            "technical_analysis": "ğŸ“ˆ æŠ€æœ¯åˆ†æ",
            "technical_report": "ğŸ“ˆ æŠ€æœ¯åˆ†æ",
            "market_analysis": "ğŸŒ å¸‚åœºåˆ†æ",
            "international_news_report": "ğŸ“° å›½é™…æ–°é—»åˆ†æ",
            "risk_analysis": "âš ï¸ é£é™©åˆ†æ",
            "valuation_analysis": "ğŸ’ ä¼°å€¼åˆ†æ",
            "valuation_report": "ğŸ’ ä¼°å€¼åˆ†æ",
            "sentiment_report": "ğŸ“Š å¸‚åœºæƒ…ç»ªåˆ†æ",
            "market_sentiment": "ğŸ­ å¸‚åœºæƒ…ç»ª",
            "investment_recommendation": "ğŸ¯ æŠ•èµ„å»ºè®®",
            "strategy_report": "â™Ÿï¸ æŠ•èµ„ç­–ç•¥æŠ¥å‘Š",
            "research_team_decision": "â™Ÿï¸ æŠ•èµ„ç­–ç•¥æŠ¥å‘Š"
        }
        
        # æŒ‰é¡ºåºæ·»åŠ æ¨¡å—
        for module_key in module_order:
            if module_key in reports:
                module_content = reports[module_key]
                if isinstance(module_content, str) and module_content.strip():
                    # ğŸ”¥ å°è¯•æ ¼å¼åŒ–JSONå†…å®¹
                    module_content = self._format_json_content(module_key, module_content)
                    
                    title = module_titles.get(module_key, module_key)
                    content_parts.append(f"## {title}")
                    content_parts.append("")
                    content_parts.append(module_content)
                    content_parts.append("")
                    content_parts.append("---")
                    content_parts.append("")
        
        # æ·»åŠ å…¶ä»–æœªåˆ—å‡ºçš„æ¨¡å—
        for module_key, module_content in reports.items():
            if module_key not in module_order:
                if isinstance(module_content, str) and module_content.strip():
                    # ğŸ”¥ å°è¯•æ ¼å¼åŒ–JSONå†…å®¹
                    module_content = self._format_json_content(module_key, module_content)
                    
                    content_parts.append(f"## {module_key}")
                    content_parts.append("")
                    content_parts.append(module_content)
                    content_parts.append("")
                    content_parts.append("---")
                    content_parts.append("")
        
        # é¡µè„š
        content_parts.append("")
        content_parts.append("---")
        content_parts.append("")
        content_parts.append("*æœ¬æŠ¥å‘Šç”± TradingAgents-CN è‡ªåŠ¨ç”Ÿæˆ*")
        content_parts.append("")
        
        markdown_content = "\n".join(content_parts)
        logger.info(f"âœ… Markdown æŠ¥å‘Šç”Ÿæˆå®Œæˆï¼Œé•¿åº¦: {len(markdown_content)} å­—ç¬¦")
        
        return markdown_content
    
    def _clean_markdown_for_pandoc(self, md_content: str) -> str:
        """æ¸…ç† Markdown å†…å®¹ï¼Œé¿å… pandoc è§£æé—®é¢˜"""
        import re

        # ç§»é™¤å¯èƒ½å¯¼è‡´ YAML è§£æé—®é¢˜çš„å†…å®¹
        # å¦‚æœå¼€å¤´æœ‰ "---"ï¼Œåœ¨å‰é¢æ·»åŠ ç©ºè¡Œ
        if md_content.strip().startswith("---"):
            md_content = "\n" + md_content

        # ğŸ”¥ ç§»é™¤å¯èƒ½å¯¼è‡´ç«–æ’çš„ HTML æ ‡ç­¾å’Œæ ·å¼
        # ç§»é™¤ writing-mode ç›¸å…³çš„æ ·å¼
        md_content = re.sub(r'<[^>]*writing-mode[^>]*>', '', md_content, flags=re.IGNORECASE)
        md_content = re.sub(r'<[^>]*text-orientation[^>]*>', '', md_content, flags=re.IGNORECASE)

        # ç§»é™¤ <div> æ ‡ç­¾ä¸­çš„ style å±æ€§ï¼ˆå¯èƒ½åŒ…å«ç«–æ’æ ·å¼ï¼‰
        md_content = re.sub(r'<div\s+style="[^"]*">', '<div>', md_content, flags=re.IGNORECASE)
        md_content = re.sub(r'<span\s+style="[^"]*">', '<span>', md_content, flags=re.IGNORECASE)

        # ğŸ”¥ ç§»é™¤å¯èƒ½å¯¼è‡´é—®é¢˜çš„ HTML æ ‡ç­¾
        # ä¿ç•™åŸºæœ¬çš„ Markdown æ ¼å¼ï¼Œç§»é™¤å¤æ‚çš„ HTML
        md_content = re.sub(r'<style[^>]*>.*?</style>', '', md_content, flags=re.DOTALL | re.IGNORECASE)

        # ğŸ”¥ ç¡®ä¿æ‰€æœ‰æ®µè½éƒ½æ˜¯æ­£å¸¸çš„æ¨ªæ’æ–‡æœ¬
        # åœ¨æ¯ä¸ªæ®µè½å‰åæ·»åŠ æ˜ç¡®çš„æ¢è¡Œï¼Œé¿å… Pandoc è¯¯åˆ¤
        lines = md_content.split('\n')
        cleaned_lines = []
        for line in lines:
            # è·³è¿‡ç©ºè¡Œ
            if not line.strip():
                cleaned_lines.append(line)
                continue

            # å¦‚æœæ˜¯æ ‡é¢˜ã€åˆ—è¡¨ã€è¡¨æ ¼ç­‰ Markdown è¯­æ³•ï¼Œä¿æŒåŸæ ·
            if line.strip().startswith(('#', '-', '*', '|', '>', '```', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                cleaned_lines.append(line)
            else:
                # æ™®é€šæ®µè½ï¼šç¡®ä¿æ²¡æœ‰ç‰¹æ®Šå­—ç¬¦å¯¼è‡´ç«–æ’
                cleaned_lines.append(line)

        md_content = '\n'.join(cleaned_lines)

        return md_content

    def _create_pdf_css(self) -> str:
        """åˆ›å»º PDF æ ·å¼è¡¨ï¼Œæ§åˆ¶è¡¨æ ¼åˆ†é¡µå’Œæ–‡æœ¬æ–¹å‘"""
        return """
<style>
/* ğŸ”¥ å¼ºåˆ¶æ‰€æœ‰æ–‡æœ¬æ¨ªæ’æ˜¾ç¤ºï¼ˆä¿®å¤ä¸­æ–‡ç«–æ’é—®é¢˜ï¼‰ */
* {
    writing-mode: horizontal-tb !important;
    text-orientation: mixed !important;
}

body {
    writing-mode: horizontal-tb !important;
    direction: ltr !important;
}

/* æ®µè½å’Œæ–‡æœ¬ */
p, div, span, td, th, li {
    writing-mode: horizontal-tb !important;
    text-orientation: mixed !important;
}

/* è¡¨æ ¼æ ·å¼ - å…è®¸è·¨é¡µ */
table {
    width: 100%;
    border-collapse: collapse;
    page-break-inside: auto;
    writing-mode: horizontal-tb !important;
}

/* è¡¨æ ¼è¡Œ - é¿å…åœ¨è¡Œä¸­é—´åˆ†é¡µ */
tr {
    page-break-inside: avoid;
    page-break-after: auto;
}

/* è¡¨å¤´ - åœ¨æ¯é¡µé‡å¤æ˜¾ç¤º */
thead {
    display: table-header-group;
}

/* è¡¨æ ¼å•å…ƒæ ¼ */
td, th {
    padding: 8px;
    border: 1px solid #ddd;
    writing-mode: horizontal-tb !important;
    text-orientation: mixed !important;
}

/* è¡¨å¤´æ ·å¼ */
th {
    background-color: #f2f2f2;
    font-weight: bold;
}

/* é¿å…æ ‡é¢˜åç«‹å³åˆ†é¡µ */
h1, h2, h3, h4, h5, h6 {
    page-break-after: avoid;
    writing-mode: horizontal-tb !important;
}

/* é¿å…åœ¨åˆ—è¡¨é¡¹ä¸­é—´åˆ†é¡µ */
li {
    page-break-inside: avoid;
}

/* ä»£ç å— */
pre, code {
    writing-mode: horizontal-tb !important;
    white-space: pre-wrap;
    word-wrap: break-word;
}
</style>
"""
    
    def generate_docx_report(self, report_doc: Dict[str, Any]) -> bytes:
        """ç”Ÿæˆ Word æ–‡æ¡£æ ¼å¼æŠ¥å‘Š"""
        logger.info("ğŸ“„ å¼€å§‹ç”Ÿæˆ Word æ–‡æ¡£...")

        if not self.pandoc_available:
            raise Exception("Pandoc ä¸å¯ç”¨ï¼Œæ— æ³•ç”Ÿæˆ Word æ–‡æ¡£ã€‚è¯·å®‰è£… pandoc æˆ–ä½¿ç”¨ Markdown æ ¼å¼å¯¼å‡ºã€‚")

        # ç”Ÿæˆ Markdown å†…å®¹
        md_content = self.generate_markdown_report(report_doc)

        try:
            # åˆ›å»ºä¸´æ—¶æ–‡ä»¶
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                output_file = tmp_file.name

            logger.info(f"ğŸ“ ä¸´æ—¶æ–‡ä»¶è·¯å¾„: {output_file}")

            # Pandoc å‚æ•°
            extra_args = [
                '--from=markdown-yaml_metadata_block',  # ç¦ç”¨ YAML å…ƒæ•°æ®å—è§£æ
                '--standalone',  # ç”Ÿæˆç‹¬ç«‹æ–‡æ¡£
                '--wrap=preserve',  # ä¿ç•™æ¢è¡Œ
                '--columns=120',  # è®¾ç½®åˆ—å®½
                '-M', 'lang=zh-CN',  # ğŸ”¥ æ˜ç¡®æŒ‡å®šè¯­è¨€ä¸ºç®€ä½“ä¸­æ–‡
                '-M', 'dir=ltr',  # ğŸ”¥ æ˜ç¡®æŒ‡å®šæ–‡æœ¬æ–¹å‘ä¸ºä»å·¦åˆ°å³
            ]

            # æ¸…ç†å†…å®¹
            cleaned_content = self._clean_markdown_for_pandoc(md_content)

            # è½¬æ¢ä¸º Word
            pypandoc.convert_text(
                cleaned_content,
                'docx',
                format='markdown',
                outputfile=output_file,
                extra_args=extra_args
            )

            logger.info("âœ… pypandoc è½¬æ¢å®Œæˆ")

            # ğŸ”¥ åå¤„ç†ï¼šä¿®å¤ Word æ–‡æ¡£ä¸­çš„æ–‡æœ¬æ–¹å‘
            try:
                from docx import Document
                doc = Document(output_file)

                # ä¿®å¤æ‰€æœ‰æ®µè½çš„æ–‡æœ¬æ–¹å‘
                for paragraph in doc.paragraphs:
                    # è®¾ç½®æ®µè½ä¸ºä»å·¦åˆ°å³
                    if paragraph._element.pPr is not None:
                        # ç§»é™¤å¯èƒ½çš„ç«–æ’è®¾ç½®
                        for child in list(paragraph._element.pPr):
                            if 'textDirection' in child.tag or 'bidi' in child.tag:
                                paragraph._element.pPr.remove(child)

                # ä¿®å¤è¡¨æ ¼ä¸­çš„æ–‡æœ¬æ–¹å‘
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if paragraph._element.pPr is not None:
                                    for child in list(paragraph._element.pPr):
                                        if 'textDirection' in child.tag or 'bidi' in child.tag:
                                            paragraph._element.pPr.remove(child)

                # ä¿å­˜ä¿®å¤åçš„æ–‡æ¡£
                doc.save(output_file)
                logger.info("âœ… Word æ–‡æ¡£æ–‡æœ¬æ–¹å‘ä¿®å¤å®Œæˆ")
            except ImportError:
                logger.warning("âš ï¸ python-docx æœªå®‰è£…ï¼Œè·³è¿‡æ–‡æœ¬æ–¹å‘ä¿®å¤")
            except Exception as e:
                logger.warning(f"âš ï¸ Word æ–‡æ¡£æ–‡æœ¬æ–¹å‘ä¿®å¤å¤±è´¥: {e}")

            # è¯»å–ç”Ÿæˆçš„æ–‡ä»¶
            with open(output_file, 'rb') as f:
                docx_content = f.read()

            logger.info(f"âœ… Word æ–‡æ¡£ç”ŸæˆæˆåŠŸï¼Œå¤§å°: {len(docx_content)} å­—èŠ‚")

            # æ¸…ç†ä¸´æ—¶æ–‡ä»¶
            os.unlink(output_file)

            return docx_content
            
        except Exception as e:
            logger.error(f"âŒ Word æ–‡æ¡£ç”Ÿæˆå¤±è´¥: {e}", exc_info=True)
            # æ¸…ç†ä¸´æ—¶æ–‡ä»¶
            try:
                if 'output_file' in locals() and os.path.exists(output_file):
                    os.unlink(output_file)
            except:
                pass
            raise Exception(f"ç”Ÿæˆ Word æ–‡æ¡£å¤±è´¥: {e}")
    
    def _markdown_to_html(self, md_content: str) -> str:
        """å°† Markdown è½¬æ¢ä¸º HTML"""
        import markdown

        # é…ç½® Markdown æ‰©å±•
        extensions = [
            'markdown.extensions.tables',  # è¡¨æ ¼æ”¯æŒ
            'markdown.extensions.fenced_code',  # ä»£ç å—æ”¯æŒ
            'markdown.extensions.nl2br',  # æ¢è¡Œæ”¯æŒ
        ]

        # è½¬æ¢ä¸º HTML
        html_content = markdown.markdown(md_content, extensions=extensions)

        # æ·»åŠ  HTML æ¨¡æ¿å’Œæ ·å¼
        # WeasyPrint ä¼˜åŒ–çš„ CSSï¼ˆç§»é™¤ä¸æ”¯æŒçš„å±æ€§ï¼‰
        html_template = f"""
<!DOCTYPE html>
<html lang="zh-CN" dir="ltr">
<head>
    <meta charset="UTF-8">
    <title>åˆ†ææŠ¥å‘Š</title>
    <style>
        /* åŸºç¡€æ ·å¼ - ç¡®ä¿æ–‡æœ¬æ–¹å‘æ­£ç¡® */
        html {{
            direction: ltr;
        }}

        body {{
            font-family: "Noto Sans CJK SC", "Microsoft YaHei", "SimHei", "Arial", sans-serif;
            line-height: 1.8;
            color: #333;
            margin: 20mm;
            padding: 0;
            background: white;
            direction: ltr;
        }}

        /* æ ‡é¢˜æ ·å¼ */
        h1, h2, h3, h4, h5, h6 {{
            color: #2c3e50;
            margin-top: 1.5em;
            margin-bottom: 0.8em;
            font-weight: 600;
            page-break-after: avoid;
            direction: ltr;
        }}

        h1 {{
            font-size: 2em;
            border-bottom: 3px solid #3498db;
            padding-bottom: 0.3em;
            page-break-before: always;
        }}

        h1:first-child {{
            page-break-before: avoid;
        }}

        h2 {{
            font-size: 1.6em;
            border-bottom: 2px solid #bdc3c7;
            padding-bottom: 0.25em;
        }}

        h3 {{
            font-size: 1.3em;
            color: #34495e;
        }}

        /* æ®µè½æ ·å¼ */
        p {{
            margin: 0.8em 0;
            text-align: left;
            direction: ltr;
        }}

        /* è¡¨æ ¼æ ·å¼ - ä¼˜åŒ–åˆ†é¡µ */
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 1.5em 0;
            font-size: 0.9em;
            direction: ltr;
        }}

        /* è¡¨å¤´åœ¨æ¯é¡µé‡å¤ */
        thead {{
            display: table-header-group;
        }}

        tbody {{
            display: table-row-group;
        }}

        /* è¡¨æ ¼è¡Œé¿å…è·¨é¡µæ–­å¼€ */
        tr {{
            page-break-inside: avoid;
        }}

        th, td {{
            border: 1px solid #ddd;
            padding: 10px 12px;
            text-align: left;
            direction: ltr;
        }}

        th {{
            background-color: #3498db;
            color: white;
            font-weight: bold;
        }}

        tbody tr:nth-child(even) {{
            background-color: #f8f9fa;
        }}

        tbody tr:hover {{
            background-color: #e9ecef;
        }}

        /* ä»£ç å—æ ·å¼ */
        code {{
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: "Consolas", "Monaco", "Courier New", monospace;
            font-size: 0.9em;
            direction: ltr;
        }}

        pre {{
            background-color: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            border-left: 4px solid #3498db;
            page-break-inside: avoid;
            direction: ltr;
        }}

        pre code {{
            background-color: transparent;
            padding: 0;
        }}

        /* åˆ—è¡¨æ ·å¼ */
        ul, ol {{
            margin: 0.8em 0;
            padding-left: 2em;
            direction: ltr;
        }}

        li {{
            margin: 0.4em 0;
            direction: ltr;
        }}

        /* å¼ºè°ƒæ–‡æœ¬ */
        strong, b {{
            font-weight: 700;
            color: #2c3e50;
        }}

        em, i {{
            font-style: italic;
            color: #555;
        }}

        /* æ°´å¹³çº¿ */
        hr {{
            border: none;
            border-top: 2px solid #ecf0f1;
            margin: 2em 0;
        }}

        /* é“¾æ¥æ ·å¼ */
        a {{
            color: #3498db;
            text-decoration: none;
        }}

        a:hover {{
            text-decoration: underline;
        }}

        /* åˆ†é¡µæ§åˆ¶ */
        @page {{
            size: A4;
            margin: 20mm;

            @top-center {{
                content: "åˆ†ææŠ¥å‘Š";
                font-size: 10pt;
                color: #999;
            }}

            @bottom-right {{
                content: "ç¬¬ " counter(page) " é¡µ";
                font-size: 10pt;
                color: #999;
            }}
        }}

        /* é¿å…å­¤è¡Œå’Œå¯¡è¡Œ */
        p, li {{
            orphans: 3;
            widows: 3;
        }}

        /* å›¾ç‰‡æ ·å¼ */
        img {{
            max-width: 100%;
            height: auto;
            page-break-inside: avoid;
        }}

        /* å¼•ç”¨å—æ ·å¼ */
        blockquote {{
            margin: 1em 0;
            padding: 0.5em 1em;
            border-left: 4px solid #3498db;
            background-color: #f8f9fa;
            font-style: italic;
            page-break-inside: avoid;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
"""
        return html_template

    def _generate_pdf_with_pdfkit(self, html_content: str) -> bytes:
        """ä½¿ç”¨ pdfkit ç”Ÿæˆ PDF"""
        import pdfkit

        logger.info("ğŸ”§ ä½¿ç”¨ pdfkit + wkhtmltopdf ç”Ÿæˆ PDF...")

        # é…ç½®é€‰é¡¹
        options = {
            'encoding': 'UTF-8',
            'enable-local-file-access': None,
            'page-size': 'A4',
            'margin-top': '20mm',
            'margin-right': '20mm',
            'margin-bottom': '20mm',
            'margin-left': '20mm',
        }

        # ç”Ÿæˆ PDF
        pdf_bytes = pdfkit.from_string(html_content, False, options=options)

        logger.info(f"âœ… pdfkit PDF ç”ŸæˆæˆåŠŸï¼Œå¤§å°: {len(pdf_bytes)} å­—èŠ‚")
        return pdf_bytes

    def _generate_pdf_with_weasyprint(self, html_content: str) -> bytes:
        """ä½¿ç”¨ weasyprint ç”Ÿæˆ PDF"""
        import weasyprint
        
        logger.info("ğŸ”§ ä½¿ç”¨ weasyprint ç”Ÿæˆ PDF...")
        
        # åˆ›å»º HTML å¯¹è±¡
        html = weasyprint.HTML(string=html_content)
        
        # æ¸²æŸ“ PDF
        pdf_bytes = html.write_pdf()
        
        logger.info(f"âœ… weasyprint PDF ç”ŸæˆæˆåŠŸï¼Œå¤§å°: {len(pdf_bytes)} å­—èŠ‚")
        return pdf_bytes

    def generate_pdf_report(self, report_doc: Dict[str, Any]) -> bytes:
        """ç”Ÿæˆ PDF æ ¼å¼æŠ¥å‘Šï¼ˆä¼˜å…ˆä½¿ç”¨ weasyprintï¼Œé™çº§ä½¿ç”¨ pdfkitï¼‰"""
        logger.info("ğŸ“Š å¼€å§‹ç”Ÿæˆ PDF æ–‡æ¡£...")

        # æ£€æŸ¥ PDF å·¥å…·æ˜¯å¦å¯ç”¨
        if not self.weasyprint_available and not self.pdfkit_available:
            error_msg = (
                "PDF ç”Ÿæˆå·¥å…·ä¸å¯ç”¨ã€‚\n\n"
                "è¯·å®‰è£…ä»¥ä¸‹ä»»ä¸€å·¥å…·:\n"
                "1. weasyprint (æ¨è): brew install pango && pip install weasyprint\n"
                "2. wkhtmltopdf: https://wkhtmltopdf.org/downloads.html\n"
            )
            if WEASYPRINT_ERROR:
                error_msg += f"\nWeasyPrint é”™è¯¯: {WEASYPRINT_ERROR}"
            if PDFKIT_ERROR:
                error_msg += f"\nPDFKit é”™è¯¯: {PDFKIT_ERROR}"

            logger.error(f"âŒ {error_msg}")
            raise Exception(error_msg)

        # ç”Ÿæˆ Markdown å†…å®¹
        md_content = self.generate_markdown_report(report_doc)
        
        try:
            # è½¬æ¢ä¸º HTML
            html_content = self._markdown_to_html(md_content)
            
            # ä¼˜å…ˆä½¿ç”¨ weasyprint
            if self.weasyprint_available:
                try:
                    return self._generate_pdf_with_weasyprint(html_content)
                except Exception as e:
                    logger.error(f"âš ï¸ weasyprint ç”Ÿæˆå¤±è´¥: {e}ï¼Œå°è¯•ä½¿ç”¨ pdfkit...")
                    if not self.pdfkit_available:
                        raise e
            
            # ä½¿ç”¨ pdfkit
            return self._generate_pdf_with_pdfkit(html_content)
            
        except Exception as e:
            error_msg = f"PDF ç”Ÿæˆå¤±è´¥: {e}"
            logger.error(f"âŒ {error_msg}")
            raise Exception(error_msg)


# åˆ›å»ºå…¨å±€å¯¼å‡ºå™¨å®ä¾‹
report_exporter = ReportExporter()

